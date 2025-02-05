import os
import io
from PIL import Image
import fitz  # PyMuPDF
import pytesseract
from PyPDF2 import PdfReader
import docx
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_google_genai import GoogleGenerativeAIEmbeddings
import streamlit as st
import google.generativeai as genai
from langchain.vectorstores import FAISS
from langchain_google_genai import ChatGoogleGenerativeAI
from langchain.chains.question_answering import load_qa_chain
from langchain.prompts import PromptTemplate
from dotenv import load_dotenv

load_dotenv()
os.getenv("GOOGLE_API_KEY")
genai.configure(api_key=os.getenv("GOOGLE_API_KEY"))

# Configure Tesseract path (update with your installation path)
pytesseract.pytesseract.tesseract_cmd = r'/usr/bin/tesseract'  # Linux/Mac
# pytesseract.pytesseract.tesseract_cmd = r'C:\Program Files\Tesseract-OCR\tesseract.exe'  # Windows

def get_document_text(uploaded_files):
    """Extract text from PDF and Word documents with OCR fallback"""
    text = ""
    for file in uploaded_files:
        if file.name.endswith('.pdf'):
            # Process PDF with OCR capability
            file_content = file.read()
            try:
                # Try normal text extraction first
                pdf_reader = PdfReader(io.BytesIO(file_content))
                for page in pdf_reader.pages:
                    page_text = page.extract_text() or ""
                    
                    # If text is insufficient, use OCR
                    if len(page_text.strip()) < 50:
                        # Use PyMuPDF for image extraction
                        doc = fitz.open(stream=file_content, filetype="pdf")
                        page_num = page.page_number
                        pix = doc.load_page(page_num).get_pixmap()
                        img = Image.open(io.BytesIO(pix.tobytes()))
                        page_text += pytesseract.image_to_string(img)
                        doc.close()
                    
                    text += page_text + "\n"
                
            except Exception as e:
                st.error(f"Error processing PDF: {str(e)}")

        elif file.name.endswith('.docx'):
            # Process Word documents
            try:
                doc = docx.Document(io.BytesIO(file.read()))
                for para in doc.paragraphs:
                    text += para.text + "\n"
                for table in doc.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            text += cell.text + "\n"
            except Exception as e:
                st.error(f"Error processing Word document: {str(e)}")
    
    return text

def get_text_chunks(text):
    """Split text into manageable chunks"""
    splitter = RecursiveCharacterTextSplitter(
        chunk_size=10000, 
        chunk_overlap=1000
    )
    chunks = splitter.split_text(text)
    return chunks

def get_vector_store(chunks):
    """Create vector embeddings store"""
    embeddings = GoogleGenerativeAIEmbeddings(
        model="models/embedding-001"
    )
    vector_store = FAISS.from_texts(chunks, embedding=embeddings)
    vector_store.save_local("faiss_index")

# Rest of the functions remain unchanged (get_conversational_chain, clear_chat_history, user_input, main)

def main():
    """Main application interface"""
    st.set_page_config(
        page_title="Smart RFP Analyser",
        page_icon="📄",
        layout="centered"
    )

    # Sidebar Configuration
    with st.sidebar:
        st.title("Control Panel")
        st.subheader("Upload Documents")
        
        uploaded_files = st.file_uploader(
            "Choose PDF or Word files",
            type=["pdf", "docx"],
            accept_multiple_files=True
        )
        
        if st.button("Process Documents"):
            if not uploaded_files:
                st.warning("Please upload files first")
                return
                
            with st.spinner("Analyzing documents..."):
                raw_text = get_document_text(uploaded_files)
                text_chunks = get_text_chunks(raw_text)
                get_vector_store(text_chunks)
                st.success("Documents processed successfully!")

        st.button("Clear Chat History", on_click=clear_chat_history)

    # Main Chat Interface
    st.title("📄 Smart RFP Analyser")
    st.caption("Upload PDF/Word documents and ask questions about their content")

    if "messages" not in st.session_state:
        clear_chat_history()

    for message in st.session_state.messages:
        with st.chat_message(message["role"]):
            st.write(message["content"])

    if prompt := st.chat_input("Ask about your documents..."):
        st.session_state.messages.append({"role": "user", "content": prompt})
        
        with st.chat_message("user"):
            st.write(prompt)
            
        with st.chat_message("assistant"):
            with st.spinner("Analyzing documents..."):
                response = user_input(prompt)
                full_response = response.get("output_text", "Could not generate response")
                st.write(full_response)
                
        st.session_state.messages.append(
            {"role": "assistant", "content": full_response}
        )

if __name__ == "__main__":
    main()
